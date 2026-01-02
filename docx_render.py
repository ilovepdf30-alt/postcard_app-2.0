from docx import Document

def _replace_in_paragraph_runs(paragraph, mapping: dict[str, str]) -> None:
    if not paragraph.runs:
        return

    full = "".join(r.text for r in paragraph.runs)
    for ph, val in mapping.items():
        if ph not in full:
            continue

        start = 0
        while True:
            idx = full.find(ph, start)
            if idx == -1:
                break
            end = idx + len(ph)

            spans = []
            pos = 0
            for ri, r in enumerate(paragraph.runs):
                rt = r.text or ""
                spans.append((ri, pos, pos + len(rt)))
                pos += len(rt)

            cover = []
            for ri, a, b in spans:
                if b <= idx:
                    continue
                if a >= end:
                    break
                cover.append((ri, a, b))
            if not cover:
                break

            first_ri, first_a, _ = cover[0]
            last_ri, last_a, _ = cover[-1]

            prefix = paragraph.runs[first_ri].text[: max(0, idx - first_a)]
            suffix = paragraph.runs[last_ri].text[max(0, end - last_a):]

            paragraph.runs[first_ri].text = prefix + val + suffix
            for ri, _, _ in cover[1:]:
                paragraph.runs[ri].text = ""

            full = "".join(r.text for r in paragraph.runs)
            start = idx + len(val)

def replace_placeholders_docx(doc: Document, mapping: dict[str, str]) -> None:
    for p in doc.paragraphs:
        _replace_in_paragraph_runs(p, mapping)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    _replace_in_paragraph_runs(p, mapping)